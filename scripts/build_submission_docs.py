from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "output" / "submission"
OUT.mkdir(parents=True, exist_ok=True)
ASSETS = ROOT / "output" / "playwright"
LOGO = ROOT / "app" / "static" / "logo-icon.png"

PURPLE = "4F46C7"
NAVY = "202044"
BLUE = "3976C8"
MUTED = "667085"
LIGHT = "EEF2FF"
PALE = "F7F8FC"


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=NAVY, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D9DDEB", size="6"):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True, "FFFFFF", 9.5)
        shade(table.rows[0].cells[i], PURPLE)
        set_cell_margins(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False, NAVY, 9.2)
            shade(cells[i], "FFFFFF" if len(table.rows) % 2 else PALE)
            set_cell_margins(cells[i])
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "C9D2FF", "8")
    cell = table.cell(0, 0)
    shade(cell, LIGHT)
    set_cell_margins(cell, 150, 170, 150, 170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label + "  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(PURPLE)
    r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor.from_string(NAVY)
    r2.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image(doc, path, width=6.2, caption=None, caption_before=False):
    if path.exists():
        if caption and caption_before:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(6)
            cp.paragraph_format.space_after = Pt(5)
            for r in cp.runs:
                r.italic = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor.from_string(MUTED)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        if caption and not caption_before:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(8)
            for r in cp.runs:
                r.italic = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor.from_string(MUTED)


def page_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def setup_doc(title):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color, before, after in (
        ("Heading 1", 17, PURPLE, 16, 8),
        ("Heading 2", 13.5, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        st = styles[style_name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.12
    header = sec.header.paragraphs[0]
    header.text = "Deepsleep深度睡眠  ·  顺德寻味 · 城市增长驾驶舱"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in header.runs:
        r.font.name = "Microsoft YaHei"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("2026 顺德黑客松  ·  ")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    page_field(footer)
    for r in footer.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(MUTED)
    return doc


def add_cover(doc, title, subtitle, document_type):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(32)
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(0.8))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run(title)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(25)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(PURPLE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run(document_type)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    add_callout(doc, "团队", "Deepsleep深度睡眠  ·  梁梓轩 / 罗其立 / 郑嘉杰 / 王顺 / 梁倪滔")
    doc.add_page_break()


def build_solution():
    doc = setup_doc("项目说明")
    add_cover(doc, "顺德寻味 · 城市增长驾驶舱", "City Flavor Radar", "项目说明与解决方案")
    doc.add_heading("一、项目摘要", level=1)
    doc.add_paragraph("顺德寻味 · 城市增长驾驶舱是一个专为顺德文旅与宣传部门打造的 AI Agent 矩阵：以社媒声呐为未来接入口，主动听见那些还没爆、但该爆的本地美食店和主理人故事，并用多智能体协作完成“发现 → 挖掘 → 诊断 → 预测 → 传播”的全链路，把“网红偶然性”变成“可计算的潜力”。当前 Demo 以真实店铺资料和结构化数据演示核心流程，后续可接入更大规模的社媒数据源。")
    add_callout(doc, "一句话", "让 AI 听见顺德城市里的弱信号，把真实小店、老手艺和主理人故事，转化为可解释、可执行、可复盘的文旅增长动作。")
    doc.add_heading("评委一分钟读懂", level=2)
    doc.add_paragraph("顺德不缺好味道，缺的是让好味道被准确发现、被耐心讲清楚、被游客真正抵达的机制。我们不想再做一个把所有店铺都变成网红的工具，而是做一张能听见城市弱信号的地图：它告诉政府该扶持谁，告诉商户为什么值得传播，也告诉城市流量来了以后如何接住。")
    add_callout(doc, "我们想留下的不是一次爆款", "而是一套能在顺德不同镇街复用的发现、策源、承接和复盘机制。")
    doc.add_heading("二、问题与机会", level=1)
    add_bullets(doc, [
        "资源很多，但政府和街区很难判断应该优先扶持哪家店。",
        "店铺有手艺和故事，但缺少能够持续生产的内容结构。",
        "曝光增长后，产能、排队、停车和食品安全风险可能被忽略。",
        "传统方案容易停留在口号，缺少指标、动作、协同对象和复盘节点。",
    ])
    doc.add_heading("三、解决方案", level=1)
    doc.add_heading("1. 发现：把顺德味道放到一张城市地图上", level=2)
    doc.add_paragraph("系统将店铺名称、品类、镇街、坐标、潜力指数和来源信息放入城市味觉地图，支持节点点击、店铺档案和潜力榜，让决策者先看见城市里分散的真实资源。")
    doc.add_heading("2. 诊断：把爆款潜力变成可解释指标", level=2)
    doc.add_paragraph("核心爆款指数由五个维度组成：人设反差度 25%、风味品相度 25%、故事叙事性 20%、情绪共鸣度 20%、素材续航度 10%。风险不是加分项，而是门槛：卫生、安全和资质问题优先处理，产能不足与主理人抗拒则进入先扶持策略。")
    add_table(doc, ["等级", "指数区间", "建议处置"], [
        ("A", "≥75", "重点孵化，进入内容策源与传播计划"),
        ("B", "60–74.9", "补齐短板，观察后再放大"),
        ("C", "<60", "不优先投入传播资源"),
    ], [1.0, 1.2, 4.2])
    doc.add_heading("3. 预测：用真实案例作为引爆力锚点", level=2)
    doc.add_paragraph("系统以莫氏鸡煲作为已验证案例锚点，将店铺指数、品类系数、素材续航和承载能力纳入估算，输出曝光、客流、排队、热度周期与变现回报。结果用于辅助决策，不冒充官方统计或已发生的商业结果。")
    doc.add_heading("4. 策源：从一家店生成一份行动方案书", level=2)
    add_bullets(doc, [
        "真实店铺画像：品类、人物、招牌、视觉画面、经营承载与风险。",
        "内容资产：15 秒、30 秒、180 秒内容脚本与官方号、达人、商户协同矩阵。",
        "90 天节奏：首发、扩散、路线承接和城市 IP 沉淀。",
        "公共服务：打卡点、导览、消费券、交通服务、排队与食品安全预案。",
    ])
    add_image(doc, ASSETS / "dashboard-regression.png", 6.25, "图 1  城市增长驾驶舱主界面与顺德味觉地图")
    doc.add_heading("四、用户与价值", level=1)
    add_table(doc, ["用户", "真实需求", "系统提供"], [
        ("文旅 / 宣传部门", "找到值得扶持的本地内容火种", "潜力分层、风险门槛、城市行动方案"),
        ("镇街 / 街区运营者", "组织路线与公共承接", "地图节点、导览、消费券、交通建议"),
        ("美食商户", "知道为什么值得传播，流量来了如何接住", "故事卡、内容脚本、承载边界与行动清单"),
        ("达人 / 高校创作者", "找到真实、可拍、可复用的题材", "店铺素材、拍摄画面、路线和话题入口"),
    ], [1.35, 2.35, 2.7])
    doc.add_heading("五、项目成果与 Demo", level=1)
    add_bullets(doc, [
        "已完成可交互 Web 驾驶舱、城市地图、店铺节点和潜力榜。",
        "已完成政府助推快速方案，默认使用本地规则引擎，可轮换 7 个真实店铺。",
        "已完成全屏方案书，按真实店铺资料展示画像、路线、内容和风险。",
        "已完成小顺文旅策划助手，使用 DeepSeek V4 Flash 进行流式对话。",
        "已完成浏览器点击回归、后端接口和移动端适配验证。",
    ])
    add_image(doc, ASSETS / "plan-book.png", 6.25, "图 2  全屏方案书：从真实店铺画像到可执行的 90 天路线")
    p = doc.add_heading("六、团队与共创愿景", level=1)
    p.paragraph_format.page_break_before = True
    doc.add_heading("为什么是我们", level=2)
    doc.add_paragraph("我们不是在远处想象顺德，而是有人从顺德出发、有人从天津带来技术、有人把商业和路演经验带进来。这让项目同时拥有本地情感、工程实现和落地表达：既愿意把家乡的味道讲好，也能把一个想法做成可演示、可解释、可继续迭代的产品。")
    doc.add_paragraph("Deepsleep深度睡眠团队由顺德本地学生、外地高校技术成员和具备商业与创客经验的成员组成。听到顺德本地黑客松后，顺德本地成员非常激动并赶回家乡参赛；团队中也有来自天津的 985 高校技术成员，以及拥有丰富活动经历、擅长商业策划与路演的伙伴。大家合作愉快，希望用技术、数据、商业和本地文化的结合，为顺德文化旅游的发展做出贡献。")
    add_table(doc, ["成员", "身份", "主要贡献"], [
        ("梁梓轩", "学生", "队长、开发"),
        ("罗其立", "学生", "PPT 制作"),
        ("郑嘉杰", "社会", "商业策划、路演"),
        ("王顺", "学生", "技术负责人"),
        ("梁倪滔", "学生", "数据分析"),
    ], [1.6, 1.2, 3.6])
    doc.add_heading("七、真实性与合规说明", level=1)
    add_bullets(doc, [
        "店铺资料来自项目数据集及 evidence 字段记录的公开来源，信息不足处标记为待核。",
        "指数、客流、曝光和变现结果属于规则 / 模型估算，不冒充官方统计或真实订单。",
        "项目最终版本未使用 Figma、Blender 或 pixelmap.amcharts.com。",
        "提交材料不包含 API Key、永久账号密码或其他敏感凭证。",
    ])
    path = OUT / "Deepsleep深度睡眠_顺德寻味·城市增长驾驶舱_项目说明.docx"
    doc.save(path)
    return path


def build_tech():
    doc = setup_doc("技术架构")
    add_cover(doc, "顺德寻味 · 城市增长驾驶舱", "Technical Architecture & Production Notes", "技术架构、工作流程与制作说明")
    doc.add_heading("一、系统架构总览", level=1)
    doc.add_paragraph("系统采用前后端分离的轻量结构：浏览器负责地图看板、方案书和交互；FastAPI 负责数据接口与 Agent 编排；规则引擎保证政府助推方案快速、可解释、可复现；DeepSeek 仅用于小顺实时对话和可选的深度 Agent 增强。")
    add_table(doc, ["层级", "组件", "作用"], [
        ("前端", "原生 HTML / CSS / JavaScript", "城市地图、节点交互、潜力榜、方案卡、全屏方案书"),
        ("服务", "FastAPI + Uvicorn", "提供 REST API、SSE 流式对话和工作流入口"),
        ("数据", "YAML / JSON", "店铺、事件、位置、诊断、预测与工作流记录"),
        ("AI", "DeepSeek V4 Flash", "小顺文旅策划助手实时回答；可选 Agent 增强"),
        ("测试", "Node / 浏览器回归", "语法检查、接口检查、真实点击和截图验证"),
    ], [1.0, 2.2, 3.2])
    doc.add_heading("二、请求与数据流", level=1)
    add_callout(doc, "快速方案", "政府助推 · 文旅增长方案默认使用本地规则引擎，不串行请求多个模型，保证现场生成速度；结果仍从对应店铺的真实结构化资料生成。")
    add_callout(doc, "实时对话", "小顺 · 文旅策划搭子通过 SSE 调用 DeepSeek V4 Flash，前端逐段显示回答；网络或模型异常时自动显示本地兜底话术。")
    doc.add_heading("三、Agent 工作流程", level=1)
    add_numbered(doc, [
        "Mission：确定城市任务、受众、预算、周期与关键指标。",
        "Radar：从事件库与店铺池中整理热点、弱信号和机会清单。",
        "Asset Card：提取店铺人物、招牌菜、视觉钩子、承载能力和风险。",
        "Spark：计算传播势能、城市承接、转化能力和风险可控性。",
        "Decode：拆解外部案例的可迁移基因，如人物、反差、视觉、事件和承接。",
        "Migrate：将外部爆点逻辑映射到目标店铺，输出稳健、突破和现象级方案。",
        "Forge：生成主题、故事线、15 / 30 / 180 秒脚本和渠道矩阵。",
        "Ops Kit / War Room：形成 90 天日历、线下承接、风险预案和复盘建议。",
    ])
    doc.add_heading("四、主要 API", level=1)
    add_table(doc, ["接口", "方法", "用途"], [
        ("/api/shops", "GET", "店铺潜力榜与基础资料"),
        ("/api/locations", "GET", "地图节点坐标与店铺关联"),
        ("/api/shops/{shop_id}", "GET", "店铺档案、诊断和预测结果"),
        ("/api/workflow/run", "POST", "生成政府助推文旅行动方案"),
        ("/api/chat", "POST / SSE", "小顺 DeepSeek 流式对话"),
        ("/api/diagnose", "POST", "单店诊断 Agent"),
        ("/api/forecast", "POST", "单店流量预测 Agent"),
    ], [2.1, 1.05, 3.25])
    doc.add_heading("五、项目文件结构", level=1)
    add_table(doc, ["目录 / 文件", "说明"], [
        ("app/main.py", "FastAPI 路由、静态页面与工作流 API"),
        ("app/config.py", "本地环境变量和 DeepSeek 配置读取"),
        ("app/llm.py", "OpenAI-compatible 客户端封装与 SSE / JSON 调用"),
        ("app/agents/", "mission、radar、asset、spark、forge、ops、warroom 等 Agent"),
        ("app/static/index.html", "单文件城市增长驾驶舱前端"),
        ("data/", "店铺、事件、坐标和来源资料"),
        ("output/", "诊断、预测、工作流记录和浏览器回归截图"),
    ], [2.2, 4.2])
    doc.add_heading("六、运行与演示", level=1)
    doc.add_paragraph("运行环境：Windows / Python 3.12+。安装 requirements.txt 后，在项目目录启动：")
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    r = p.add_run("python -m uvicorn app.main:app --host 127.0.0.1 --port 8000")
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(PURPLE)
    add_numbered(doc, [
        "打开首页，查看顺德味觉地图和店铺节点。",
        "点击店铺，阅读店铺档案、潜力指数和风险。",
        "点击政府助推按钮，快速生成对应店铺方案。",
        "打开完整方案书，查看画像、90 天路线、内容和承载。",
        "进入小顺助手，现场追问路线、内容或风险问题。",
    ])
    doc.add_heading("七、验证与故障兜底", level=1)
    add_bullets(doc, [
        "Node 语法检查：前端脚本通过。",
        "FastAPI 导入检查：服务路由加载正常。",
        "接口检查：主页、地图资源、店铺、方案和聊天接口均已验证。",
        "浏览器回归：地图节点点击、潜力榜、方案书目录、弹层关闭、移动端布局均已点击验证。",
        "规则兜底：政府助推方案不依赖网络；小顺聊天在 DeepSeek 不可用时不会阻塞页面。",
    ])
    add_image(doc, ASSETS / "solution-card.png", 4.7, "图 1  政府助推快速方案卡", caption_before=True)
    add_image(doc, ASSETS / "plan-book.png", 4.7, "图 2  完整方案书阅读界面", caption_before=True)
    doc.add_heading("八、AI 工具、素材与真实性承诺", level=1)
    doc.add_paragraph("项目使用 DeepSeek V4 Flash 进行小顺流式对话；使用本地规则引擎完成快速方案、指数计算与风险校验；使用 FastAPI、Uvicorn、PyYAML、Pydantic 和 OpenAI Python Client 等开源组件。项目最终版本未使用 Figma、Blender 或 pixelmap.amcharts.com。地图、Logo、小顺 IP 和店铺资料均作为项目本地资源管理；公开数据来源保留在数据集 evidence 字段中。")
    doc.add_paragraph("团队承诺不把模型生成内容冒充官方统计、实地采访、真实订单或已发生的商业结果；提交材料不包含 API Key、永久账号密码或其他敏感凭证。")
    path = OUT / "Deepsleep深度睡眠_顺德寻味·城市增长驾驶舱_技术架构与制作说明.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build_solution())
    print(build_tech())
